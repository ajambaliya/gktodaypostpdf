import io
import os
import requests
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime
import pymongo
from deep_translator import GoogleTranslator, exceptions
import time
import asyncio
import telegram
import logging
import tempfile
import subprocess

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# MongoDB setup
DB_NAME = os.environ.get('DB_NAME')
COLLECTION_NAME = os.environ.get('COLLECTION_NAME')
MONGO_CONNECTION_STRING = os.environ.get('MONGO_CONNECTION_STRING')

if not all([DB_NAME, COLLECTION_NAME, MONGO_CONNECTION_STRING]):
    raise ValueError("One or more required MongoDB environment variables are not set")

client = pymongo.MongoClient(MONGO_CONNECTION_STRING)
db = client[DB_NAME]
collection = db[COLLECTION_NAME]

def fetch_article_urls(base_url, pages):
    article_urls = []
    for page in range(1, pages + 1):
        url = base_url if page == 1 else f"{base_url}page/{page}/"
        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')
        for h1_tag in soup.find_all('h1', id='list'):
            a_tag = h1_tag.find('a')
            if a_tag and a_tag.get('href'):
                article_urls.append(a_tag['href'])
    return article_urls

def translate_to_gujarati(text):
    max_retries = 3
    for attempt in range(max_retries):
        try:
            translator = GoogleTranslator(source='auto', target='gu')
            return translator.translate(text)
        except exceptions.TranslationNotFoundException as e:
            logging.warning(f"Translation not found: {e}")
            return text
        except Exception as e:
            logging.error(f"Error in translation (attempt {attempt + 1}): {e}")
            time.sleep(2)
    return text

async def scrape_and_get_content(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    main_content = soup.find('div', class_='inside_post column content_width')
    if not main_content:
        raise Exception("Main content div not found")
    heading = main_content.find('h1', id='list')
    if not heading:
        raise Exception("Heading not found")
    content_list = []
    heading_text = heading.get_text()
    translated_heading = translate_to_gujarati(heading_text)
    content_list.append({'type': 'heading', 'text': translated_heading})
    content_list.append({'type': 'heading', 'text': heading_text})
    for tag in main_content.find_all(recursive=False):
        if tag.get('class') in [['sharethis-inline-share-buttons', 'st-center', 'st-has-labels', 'st-inline-share-buttons', 'st-animated'], ['prenext']]:
            continue
        text = tag.get_text()
        translated_text = translate_to_gujarati(text)
        if tag.name == 'p':
            content_list.append({'type': 'paragraph', 'text': translated_text})
            content_list.append({'type': 'paragraph', 'text': text})
        elif tag.name == 'h2':
            content_list.append({'type': 'heading_2', 'text': translated_text})
            content_list.append({'type': 'heading_2', 'text': text})
        elif tag.name == 'h4':
            content_list.append({'type': 'heading_4', 'text': translated_text})
            content_list.append({'type': 'heading_4', 'text': text})
        elif tag.name == 'ul':
            for li in tag.find_all('li'):
                li_text = li.get_text()
                translated_li_text = translate_to_gujarati(li_text)
                content_list.append({'type': 'list_item', 'text': f"• {translated_li_text}"})
                content_list.append({'type': 'list_item', 'text': f"• {li_text}"})
    return content_list

def insert_content_between_placeholders(doc, content_list):
    start_placeholder = None
    end_placeholder = None
    
    for i, para in enumerate(doc.paragraphs):
        if "START_CONTENT" in para.text:
            start_placeholder = i
        elif "END_CONTENT" in para.text:
            end_placeholder = i
            break
    
    if start_placeholder is None or end_placeholder is None:
        logging.error("Error: Could not find both placeholders")
        return

    for i in range(end_placeholder - 1, start_placeholder, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)

    content_list = content_list[::-1]

    for content in content_list:
        if content['type'] == 'heading':
            doc.paragraphs[start_placeholder]._element.addnext(doc.add_heading(content['text'], level=1)._element)
        elif content['type'] == 'paragraph':
            doc.paragraphs[start_placeholder]._element.addnext(doc.add_paragraph(content['text'], style='Normal')._element)
        elif content['type'] == 'heading_2':
            doc.paragraphs[start_placeholder]._element.addnext(doc.add_heading(content['text'], level=2)._element)
        elif content['type'] == 'heading_4':
            doc.paragraphs[start_placeholder]._element.addnext(doc.add_heading(content['text'], level=4)._element)
        elif content['type'] == 'list_item':
            doc.paragraphs[start_placeholder]._element.addnext(doc.add_paragraph(content['text'], style='List Bullet')._element)

    doc.paragraphs[start_placeholder].text = ""
    doc.paragraphs[end_placeholder].text = ""

def download_template(url):
    download_url = url.replace('/edit?usp=sharing', '/export?format=docx')
    try:
        response = requests.get(download_url)
        response.raise_for_status()
        return io.BytesIO(response.content)
    except requests.exceptions.RequestException as e:
        logging.error(f"Failed to download template: {e}")
        raise

def check_and_insert_urls(urls):
    new_urls = []
    for url in urls:
        if not collection.find_one({'url': url}):
            new_urls.append(url)
            collection.insert_one({'url': url})
    return new_urls

def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', 
                        os.path.dirname(pdf_path), docx_path], 
                       check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        
        # Rename the output file to match the desired pdf_path
        original_pdf = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
        original_pdf_path = os.path.join(os.path.dirname(pdf_path), original_pdf)
        os.rename(original_pdf_path, pdf_path)
        
        logging.info(f"Converted {docx_path} to {pdf_path} successfully")
    except subprocess.CalledProcessError as e:
        logging.error(f"Failed to convert DOCX to PDF: {e}")
        logging.error(f"STDOUT: {e.stdout.decode()}")
        logging.error(f"STDERR: {e.stderr.decode()}")
        raise

async def send_pdf_to_telegram(pdf_path, bot_token, channel_id):
    bot = telegram.Bot(token=bot_token)
    max_retries = 3
    
    current_date = datetime.now().strftime("%d %B %Y")
    file_name = f"{current_date} Current Affairs.pdf"
    
    for attempt in range(max_retries):
        try:
            with open(pdf_path, 'rb') as pdf_file:
                await bot.send_document(chat_id=channel_id, document=pdf_file, filename=file_name)
            logging.info(f"PDF sent to Telegram channel as '{file_name}'")
            break
        except telegram.error.TimedOut as e:
            logging.warning(f"Timeout error (attempt {attempt + 1}): {e}")
            time.sleep(5)
        except Exception as e:
            logging.error(f"Error sending document (attempt {attempt + 1}): {e}")
            break

async def main():
    try:
        # Debug logging for environment variables
        logging.info(f"DB_NAME: {os.environ.get('DB_NAME')}")
        logging.info(f"COLLECTION_NAME: {os.environ.get('COLLECTION_NAME')}")
        logging.info(f"MONGO_CONNECTION_STRING: {'*' * len(os.environ.get('MONGO_CONNECTION_STRING', '')) if os.environ.get('MONGO_CONNECTION_STRING') else 'Not set'}")
        logging.info(f"TEMPLATE_URL: {os.environ.get('TEMPLATE_URL')}")
        logging.info(f"TELEGRAM_BOT_TOKEN: {'*' * len(os.environ.get('TELEGRAM_BOT_TOKEN', '')) if os.environ.get('TELEGRAM_BOT_TOKEN') else 'Not set'}")
        logging.info(f"TELEGRAM_CHANNEL_ID: {os.environ.get('TELEGRAM_CHANNEL_ID')}")

        base_url = "https://www.gktoday.in/current-affairs/"
        article_urls = fetch_article_urls(base_url, 2)
        new_urls = check_and_insert_urls(article_urls)
        if not new_urls:
            logging.info("No new articles found to scrape.")
            return
        
        template_url = os.environ.get('TEMPLATE_URL')
        if not template_url:
            raise ValueError("TEMPLATE_URL environment variable is not set")
        
        template_bytes = download_template(template_url)
        
        doc = Document(template_bytes)
        logging.info("Template loaded successfully")
        
        all_content = []
        for url in new_urls:
            logging.info(f"Scraping: {url}")
            content_list = await scrape_and_get_content(url)
            all_content.extend(content_list)
        
        insert_content_between_placeholders(doc, all_content)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
            doc.save(tmp_docx.name)
        
        pdf_path = tmp_docx.name.replace('.docx', '.pdf')
        
        convert_docx_to_pdf(tmp_docx.name, pdf_path)
        
        bot_token = os.environ.get('TELEGRAM_BOT_TOKEN')
        channel_id = os.environ.get('TELEGRAM_CHANNEL_ID')
        
        if not bot_token or not channel_id:
            raise ValueError("TELEGRAM_BOT_TOKEN or TELEGRAM_CHANNEL_ID environment variable is not set")
        
        await send_pdf_to_telegram(pdf_path, bot_token, channel_id)
        
        os.unlink(tmp_docx.name)
        os.unlink(pdf_path)
        
    except Exception as e:
        logging.error(f"An error occurred: {e}", exc_info=True)

if __name__ == "__main__":
    asyncio.run(main())
